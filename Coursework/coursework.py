'''
importing all the needed libraries 
'''

import re
import tkinter as tk
import tkinter.ttk as ttk
import os.path
import sqlite3
import smtplib
import datetime
import docx
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
from datetime import date
from dateutil import parser
from tkcalendar import Calendar, DateEntry
from PIL import Image, ImageTk
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from tkinter import messagebox
from tkinter import filedialog
from cryptography.fernet import Fernet
import inspect, os
import shutil
'''
This class is responsible for the login functionality
'''
class Login:
    def __init__(self, openLogin = None):
        self.db_name = "dentistry.db"
        self.error = False
        self.loginScreen = tk.Tk()
        loginScreen = self.loginScreen
        self.loginScreen.geometry("405x250+491+216")
        self.loginScreen.title("SMILEXE")
        self.loginScreen.configure(background="#ffffff")
        self.loginScreen.resizable(0, 0)
        StaffIDLabel = tk.Label(self.loginScreen,text="StaffID",background="#ffffff")
        self.StaffIDEntry = tk.Entry(self.loginScreen)
        PasswordLabel = tk.Label(self.loginScreen,text="Password",background="#ffffff")
        self.PasswordEntry = tk.Entry(self.loginScreen)
        self.PasswordEntry.configure(show="*")
        quitButton = tk.Button(self.loginScreen, text="quit",background='#33C6DE', command=lambda:self.quitAndLogout(loginScreen,"Quit",None))
        loginButton = tk.Button(self.loginScreen, text="login",background='#33C6DE', command=lambda:self.login())

        if openLogin == True:
            path = "smileExe2.png"
            image = Image.open(path)
            image = image.resize((200, 300), Image.ANTIALIAS)
            self.img = ImageTk.PhotoImage(Image.open(path))
            imageLabel = tk.Label(self.loginScreen, image = self.img)
            imageLabel.image = self.img
            imageLabel2 = tk.Label(self.loginScreen, image = self.img)
            imageLabel2.image = self.img
            imageLabel.configure(background='#ffffff')
            imageLabel2.configure(background='#ffffff')
            imageLabel.grid(rowspan=6,column=1,row=0)
            imageLabel2.grid(rowspan=6,column=5,row=0)
        
        StaffIDLabel.grid(row=1,column=3,padx=10)
        self.StaffIDEntry.grid(row=2,column=3,padx=10)
        PasswordLabel.grid(row=3,column=3,padx=10)
        self.PasswordEntry.grid(row=4,column=3,padx=10)
        loginButton.grid(row=5,column=3,pady=5,padx=5)
        quitButton.grid(row=6,column=3,pady=5,padx=5)

        self.backup(False)

        self.key = b'xysnaKSNQxqiBaPss0d5GFgOLJ2Uvrbq9bs36852GJQ=' # Use one of the methods to get a key (it must be the same when decrypting)
    '''
    This will encypt a file. 
    '''
    def encrypt(self,inputFile,outputFile):
        try:
            with open(inputFile, 'rb') as f:
                data = f.read()
            fernet = Fernet(self.key)
            encrypted = fernet.encrypt(data)
            with open(outputFile, 'wb') as f:
                f.write(encrypted)
            os.remove(inputFile)
        except FileNotFoundError:
            pass
    '''
    This will decypt a file. 
    '''
    def decrypt(self,outputFile,inputFile):
        try:
            with open(inputFile, 'rb') as f:
                data = f.read()
            fernet = Fernet(self.key)
            encrypted = fernet.decrypt(data)
            with open(outputFile, 'wb') as f:
                f.write(encrypted)
            os.remove(inputFile)
        except FileNotFoundError:
            pass
    '''
    This will destroy all windows add to the log file that the user has logged out.
    It will open the login window depending on what button the user has entered.
    '''
    def quitAndLogout(self,window,choice,StaffID):
        loggingOut = False
        if messagebox.askyesno("Verify", "Are you sure you want to {0}?".format(choice)):
            db_name = self.db_name
            encryptedFileName = 'dentistry.encrypted'
            self.encrypt(db_name,encryptedFileName)
            window.destroy()
            loggingOut = True
            try:
                try:
                    self.root.destroy()
                except tk.TclError:
                    pass
            except AttributeError:
                pass
            if StaffID != None:
                today = datetime.datetime.now()
                log = "StaffID: {0}, logged out at {1}".format(StaffID, today)
                file = open('logs.txt', 'a')
                file.write(log)
                file.write("\n")
                file.close()
        if choice == "Logout" and loggingOut == True:
            Login(True)

    '''
    Unpacks each element within the list and saves each of the elements into a new list
    '''
    def unpackList(self,listToUnpack):
        tempList = []
        for i in range(len(listToUnpack)):
            data, = listToUnpack[i]
            str(data)
            tempList.append(data)
        return tempList
    '''
    below is all the validation routines used
    '''
    def presenceCheck(self,data,fieldName):
        if data == "":
            messagebox.showerror("Error", "Something must be entered into the {0} field.".format(fieldName))
            self.error = True

    def rangeCheck(self,data,uBoundry,lBoundry,fieldName):
        try:
            if int(data) > uBoundry or int(data) < lBoundry:
                messagebox.showerror("Error", "The field {0} must be within the range {1}-{2}.".format(fieldName,lBoundry,uBoundry))
                self.error = True
        except ValueError:
            messagebox.showerror("Error", "The field {0} must be within the range {1}-{2}.".format(fieldName,lBoundry,uBoundry))
            self.error = True

    def lengthCheck(self,data,uBoundry,lBoundry,fieldName):
        if len(data) > uBoundry or len(data) < lBoundry:
            messagebox.showerror("Error", "The field {0} must be within the length {1}-{2}.".format(fieldName,lBoundry,uBoundry))
            self.error = True

    def typeCheck(self,data,dataType,fieldName):
        dataType = "<class '{0}'>".format(dataType)
        try:
            testing = str(type(float(data)))
        except ValueError:
            testing = str(type(data))
        if dataType != testing:
            messagebox.showerror("Error","You must enter a {0} for the field {1}.".format(dataType,fieldName))
            self.error = True

    def formatCheck(self,data,Format,fieldName,formatNeeded):
        if re.match(r"{0}".format(Format),data):
            pass
        else:
            messagebox.showerror("Error","You must enter {0} for the field {1}.".format(formatNeeded,fieldName))
            self.error = True
    '''
    This will backup the database periodically or whenever a button is pressed. It will back the file up in a folder called BACKUP and will name the file dentistryV{number of files witin the BACKUP folder}
    '''
    def backup(self, buttonPressed):
        paths = str(os.path.dirname(os.path.abspath(inspect.getfile(inspect.currentframe())))) # script directory
        path = paths + "\BACKUP"
        listOfFiles = os.listdir(path)
        fileCount = len(listOfFiles) 
        fileCount = fileCount + 1
        if buttonPressed == False:
            today = datetime.datetime.now()
            Today = today.strftime("%d/%m/%Y")
            file = open('backupMadeLog.txt', 'r')
            lines = file.read().splitlines()
            lastLine = lines[-1]
            file.close()
            month = today.strftime("%m")
            lastMonth = int(month) - 1
            lastMonthDate = today.replace(month=lastMonth)
            lastMonthDate = lastMonthDate.strftime("%d/%m/%Y")
            if str(lastLine) == str(lastMonthDate):
                makeBackup = True
                file = open('backupMadeLog.txt', 'a')
                backupDate = "{0}".format(Today)
                file.write("\n")
                file.write(backupDate)
                file.close()
            else:
                makeBackup = False
        else:
            makeBackup = True
        if makeBackup == True:
            path = path + "\dentistryBackupV{0}.db".format(fileCount)
            shutil.copy(self.db_name, path)
            messagebox.showinfo("Backup","Backup of database has been created.")
    '''
    This will add into a text file the time and date a user has logged into the system.
    '''
    def logging(self,StaffID):
        today = datetime.datetime.now()
        log = "StaffID: {0}, logged in at {1}".format(StaffID, today)
        file = open('logs.txt', 'a')
        file.write(log)
        file.write("\n")
        file.close()
    '''
    This runs whenever the user presses the login button it is responsible for validating a users login.
    '''
    def login(self):
        db_name = self.db_name
        encryptedFileName = 'dentistry.encrypted'
        self.decrypt(db_name,encryptedFileName)
        enteredStaffID = str(self.StaffIDEntry.get())
        enteredPassword = str(self.PasswordEntry.get())
        conn = sqlite3.connect(self.db_name)
        cur = conn.cursor()
        cur.execute("select StaffID from Staff")
        StaffIDsFromDB = cur.fetchall()
        listOfStaffIDs = self.unpackList(StaffIDsFromDB)
        cur.execute("select Password from Staff")
        passwordsFromDB = cur.fetchall()
        conn.close()
        listOfPasswords = self.unpackList(passwordsFromDB)
        loggedIn = False
        for i in range(len(listOfStaffIDs)):
            if enteredStaffID == str(listOfStaffIDs[i]) and enteredPassword == str(listOfPasswords[i]):
                self.loginScreen.destroy()
                loggedIn = True
                break
        if loggedIn == True:
            db_name = self.db_name
            encryptedFileName = 'dentistry.encrypted'
            self.decrypt(db_name,encryptedFileName)
            conn = sqlite3.connect(self.db_name)
            cur = conn.cursor()
            speechmark = '"'
            StaffID = speechmark + str(listOfStaffIDs[i]) + speechmark
            cur.execute("select accessLevel from Staff where StaffID == {0}".format(StaffID))
            self.accessLevel, = cur.fetchone()
            conn.close()
            self.logging(StaffID)
            window = tk.Tk()
            mainScreen(window,self.accessLevel,StaffID)
        else:
            db_name = self.db_name
            encryptedFileName = 'dentistry.encrypted'
            self.encrypt(db_name,encryptedFileName)
            messagebox.showerror("Error","You have entered a incorrect password or username.")
            
"""
This creates a custom ttk Notebook which has a close button on every tab
"""
class CustomNotebook(ttk.Notebook):

    __initialized = False

    def __init__(self, *args, **kwargs):
        if not self.__initialized:
            self.__initialize_custom_style()
            self.__inititialized = True

        kwargs["style"] = "CustomNotebook"
        ttk.Notebook.__init__(self, *args, **kwargs)

        self._active = None

        self.bind("<ButtonPress-1>", self.on_close_press, True)
        self.bind("<ButtonRelease-1>", self.on_close_release)

    """
    Called when the button is pressed over the close button
    """
    def on_close_press(self, event):

        element = self.identify(event.x, event.y)

        if "close" in element:
            index = self.index("@%d,%d" % (event.x, event.y))
            self.state(['pressed'])
            self._active = index

    """
    Called when the button is released over the close button
    """
    def on_close_release(self, event):
        if not self.instate(['pressed']):
            return

        element =  self.identify(event.x, event.y)
        try:
            index = self.index("@%d,%d" % (event.x, event.y))
        except tk.TclError:
            pass

        if "close" in element and self._active == index:
            self.forget(index)
            self.event_generate("<<NotebookTabClosed>>")

        self.state(["!pressed"])
        self._active = None
    """
    This creates the style for the notebook placing the image for the close to the right of the name of the tab
    """
    def __initialize_custom_style(self):
        style = ttk.Style()
        self.images = (
            tk.PhotoImage("img_close", data='''
                R0lGODlhCAAIAMIBAAAAADs7O4+Pj9nZ2Ts7Ozs7Ozs7Ozs7OyH+EUNyZWF0ZWQg
                d2l0aCBHSU1QACH5BAEKAAQALAAAAAAIAAgAAAMVGDBEA0qNJyGw7AmxmuaZhWEU
                5kEJADs=
                '''),
            tk.PhotoImage("img_closeactive", data='''
                R0lGODlhCAAIAMIEAAAAAP/SAP/bNNnZ2cbGxsbGxsbGxsbGxiH5BAEKAAQALAAA
                AAAIAAgAAAMVGDBEA0qNJyGw7AmxmuaZhWEU5kEJADs=
                '''),
            tk.PhotoImage("img_closepressed", data='''
                R0lGODlhCAAIAMIEAAAAAOUqKv9mZtnZ2Ts7Ozs7Ozs7Ozs7OyH+EUNyZWF0ZWQg
                d2l0aCBHSU1QACH5BAEKAAQALAAAAAAIAAgAAAMVGDBEA0qNJyGw7AmxmuaZhWEU
                5kEJADs=
            ''')
        )

        style.element_create("close", "image", "img_close",
                            ("active", "pressed", "!disabled", "img_closepressed"),
                            ("active", "!disabled", "img_closeactive"), border=8, sticky='')
        style.layout("CustomNotebook", [("CustomNotebook.client", {"sticky": "nswe"})])
        style.layout("CustomNotebook.Tab", [
            ("CustomNotebook.tab", {
                "sticky": "nswe", 
                "children": [
                    ("CustomNotebook.padding", {
                        "side": "top", 
                        "sticky": "nswe",
                        "children": [
                            ("CustomNotebook.focus", {
                                "side": "top", 
                                "sticky": "nswe",
                                "children": [
                                    ("CustomNotebook.label", {"side": "left", "sticky": ''}),
                                    ("CustomNotebook.close", {"side": "left", "sticky": ''}),
                                ]
                        })
                    ]
                })
            ]
        })
    ])


'''
this is a class which will create the GUI for the main menu and creates the treeview
'''
class mainScreen(Login):
    def __init__(self, top=None, accessLevel=None, StaffID=None):
        super().__init__(False)
        self.loginScreen.destroy()
        self.accessLevel = accessLevel
        self.StaffID = StaffID
        self.currentTableName = None
        '''This class configures and populates the toplevel window.
           top is ,the toplevel containing window.'''
        _bgcolor = '#ffffff'  # X11 color: 'white'
        _fgcolor = '#000000'  # X11 color: 'black'
        _compcolor = '#ffffff' # X11 color: 'white'
        _ana1color = '#ffffff' # X11 color: 'gray85'
        _ana2color = '#ececec' # Closest X11 color: 'gray92'
        font11 = "-family {Segoe UI} -size 10 -weight normal -slant "  \
            "roman -underline 0 -overstrike 0"
        font9 = "-family {Times} -size 36 -weight normal -slant "  \
            "roman -underline 1 -overstrike 0"
        self.style = ttk.Style()
        self.style.theme_use('winnative')
        self.style.configure('.',background=_bgcolor)
        self.style.configure('.',foreground=_fgcolor)
        self.style.configure('.',font="TkDefaultFont")
        self.style.map('.',background=[('selected', _compcolor), ('active',_ana2color)])

        top.geometry("830x509+491+216")
        top.title("SMILEXE")
        top.configure(background="#ffffff")
        top.resizable(0, 0)
        self.top = top

        '''
        creation of the frame which the treeview will be stored in
        '''
        self.Frame1 = tk.Frame(self.top)
        self.Frame1.place(relx=0.243, rely=0.275, relheight=0.645, relwidth=0.75)
        self.Frame1.configure(relief='groove')
        self.Frame1.configure(borderwidth="2")
        self.Frame1.configure(relief="groove")
        self.Frame1.configure(background="#ffffff")

        self.nb=CustomNotebook(self.Frame1)

        ''''
        below is the creation of the buttons, entry widgets and labels for the main menu
        '''

        path = "smileExe2.png"
        image = Image.open(path)
        image = image.resize((200, 300), Image.ANTIALIAS)
        img = ImageTk.PhotoImage(Image.open(path))
        panel = tk.Label(top, image = img)
        panel.image = img
        panel.place(relx=0.012, rely=0.03, height=115, width=183)
        panel.configure(background="#ffffff")

        self.Label1 = tk.Label(self.top)
        self.Label1.place(relx=0.277, rely=0.02, height=71, width=544)
        self.Label1.configure(background="#ffffff")
        self.Label1.configure(disabledforeground="#a3a3a3")
        self.Label1.configure(font=font9)
        self.Label1.configure(foreground="#33C6DE")#3FA9C0(other shade of blue)
        self.Label1.configure(text='''Database Management''')

        self.Label2 = tk.Label(top)
        self.Label2.place(relx=0.22, rely=0.165, height=20, width=200)
        self.Label2.configure(background="#ffffff")
        self.Label2.configure(disabledforeground="#a3a3a3")
        self.Label2.configure(font=font11)
        self.Label2.configure(foreground="#000000")
        self.Label2.configure(text='''Select the field to search by: ''')

        self.AddButton = tk.Button(self.top)
        self.AddButton.place(relx=0.775, rely=0.177, height=30, width=55)
        self.AddButton.configure(activebackground="#ececec")
        self.AddButton.configure(activeforeground="#000000")
        self.AddButton.configure(background="#33C6DE")
        self.AddButton.configure(disabledforeground="#a3a3a3")
        self.AddButton.configure(foreground="#000000")
        self.AddButton.configure(highlightbackground="#33C6DE")
        self.AddButton.configure(highlightcolor="black")
        self.AddButton.configure(pady="0")
        self.AddButton.configure(text='''Add''')
        self.AddButton.configure(command=lambda:self.checkButtonPressed(1))

        self.EditButton = tk.Button(self.top)
        self.EditButton.place(relx=0.922, rely=0.177, height=30, width=55)
        self.EditButton.configure(activebackground="#ececec")
        self.EditButton.configure(activeforeground="#000000")
        self.EditButton.configure(background="#33C6DE")
        self.EditButton.configure(disabledforeground="#a3a3a3")
        self.EditButton.configure(foreground="#000000")
        self.EditButton.configure(highlightbackground="#33C6DE")
        self.EditButton.configure(highlightcolor="black")
        self.EditButton.configure(pady="0")
        self.EditButton.configure(text='''Edit''')
        self.EditButton.configure(command=lambda:self.checkButtonPressed(2))

        self.DeleteButton = tk.Button(self.top)
        self.DeleteButton.place(relx=0.848, rely=0.177, height=30, width=55)
        self.DeleteButton.configure(activebackground="#ececec")
        self.DeleteButton.configure(activeforeground="#000000")
        self.DeleteButton.configure(background="#33C6DE")
        self.DeleteButton.configure(disabledforeground="#a3a3a3")
        self.DeleteButton.configure(foreground="#000000")
        self.DeleteButton.configure(highlightbackground="#33C6DE")
        self.DeleteButton.configure(highlightcolor="black")
        self.DeleteButton.configure(pady="0")
        self.DeleteButton.configure(text='''Delete''')
        self.DeleteButton.configure(command=lambda:self.checkButtonPressed(3))

        self.clearSearchButton = tk.Button(self.top)
        self.clearSearchButton.place(relx=0.701, rely=0.177, height=30, width=55)
        self.clearSearchButton.configure(activebackground="#ececec")
        self.clearSearchButton.configure(activeforeground="#000000")
        self.clearSearchButton.configure(background="#33C6DE")
        self.clearSearchButton.configure(disabledforeground="#a3a3a3")
        self.clearSearchButton.configure(foreground="#000000")
        self.clearSearchButton.configure(highlightbackground="#33C6DE")
        self.clearSearchButton.configure(highlightcolor="black")
        self.clearSearchButton.configure(pady="0")
        self.clearSearchButton.configure(text='''Reset''')
        self.clearSearchButton.configure(command=lambda:self.refresh())

        self.SearchEntry = tk.Entry(self.top)
        self.SearchEntry.place(relx=0.469, rely=0.177, height=30, relwidth=0.19)
        self.SearchEntry.configure(background="#ffffff")
        self.SearchEntry.configure(borderwidth="3")
        self.SearchEntry.configure(disabledforeground="#a3a3a3")
        self.SearchEntry.configure(font="TkFixedFont")
        self.SearchEntry.configure(foreground="#000000")
        self.SearchEntry.configure(insertbackground="black")

        searchIcon = tk.PhotoImage(file="searchIcon.png")
        searchIcon = searchIcon.subsample(4, 4) 

        self.SearchButton = tk.Button(top)
        self.SearchButton.place(relx=0.659, rely=0.177, height=30, width=33)
        self.SearchButton.configure(activebackground="#ececec")
        self.SearchButton.configure(activeforeground="#000000")
        self.SearchButton.configure(background="#ffffff")
        self.SearchButton.configure(disabledforeground="#a3a3a3")
        self.SearchButton.configure(foreground="#000000")
        self.SearchButton.configure(highlightbackground="#32C1DC")
        self.SearchButton.configure(highlightcolor="black")
        self.SearchButton.configure(pady="0")
        self.SearchButton.configure(borderwidth="0")
        self.SearchButton.configure(command=lambda:self.checkButtonPressed(4))
        self.SearchButton.configure(image = searchIcon, compound=tk.RIGHT)
        self.SearchButton.image = searchIcon
    
        self.dropDownMenu(None)

        self.reportButton = tk.Button(self.top)
        self.reportButton.place(relx=0.67, rely=0.925, height=30, width=57)
        self.reportButton.configure(activebackground="#ececec")
        self.reportButton.configure(activeforeground="#000000")
        self.reportButton.configure(background="#32C1DC")
        self.reportButton.configure(disabledforeground="#a3a3a3")
        self.reportButton.configure(foreground="#000000")
        self.reportButton.configure(highlightbackground="#32C1DC")
        self.reportButton.configure(highlightcolor="black")
        self.reportButton.configure(pady="0")
        self.reportButton.configure(text='''report''')
        self.reportButton.configure(command=lambda:self.report())
        
        self.backupButton = tk.Button(self.top)
        self.backupButton.place(relx=0.75, rely=0.925, height=30, width=57)
        self.backupButton.configure(activebackground="#ececec")
        self.backupButton.configure(activeforeground="#000000")
        self.backupButton.configure(background="#32C1DC")
        self.backupButton.configure(disabledforeground="#a3a3a3")
        self.backupButton.configure(foreground="#000000")
        self.backupButton.configure(highlightbackground="#32C1DC")
        self.backupButton.configure(highlightcolor="black")
        self.backupButton.configure(pady="0")
        self.backupButton.configure(text='''backup''')
        self.backupButton.configure(command=lambda:self.backup(True))
        
        self.QuitButton = tk.Button(self.top)
        self.QuitButton.place(relx=0.91, rely=0.925, height=30, width=57)
        self.QuitButton.configure(activebackground="#ececec")
        self.QuitButton.configure(activeforeground="#000000")
        self.QuitButton.configure(background="#32C1DC")
        self.QuitButton.configure(disabledforeground="#a3a3a3")
        self.QuitButton.configure(foreground="#000000")
        self.QuitButton.configure(highlightbackground="#32C1DC")
        self.QuitButton.configure(highlightcolor="black")
        self.QuitButton.configure(pady="0")
        self.QuitButton.configure(text='''Quit''')
        self.QuitButton.configure(command=lambda:self.quitAndLogout(top,"Quit",StaffID))

        self.LogoutButton = tk.Button(top)
        self.LogoutButton.place(relx=0.83, rely=0.925, height=30, width=57)
        self.LogoutButton.configure(activebackground="#ececec")
        self.LogoutButton.configure(activeforeground="#000000")
        self.LogoutButton.configure(background="#32C1DC")
        self.LogoutButton.configure(disabledforeground="#a3a3a3")
        self.LogoutButton.configure(foreground="#000000")
        self.LogoutButton.configure(highlightbackground="#32C1DC")
        self.LogoutButton.configure(highlightcolor="black")
        self.LogoutButton.configure(pady="0")
        self.LogoutButton.configure(text='''Logout''')
        self.LogoutButton.configure(command=lambda:self.quitAndLogout(top,"Logout",StaffID))

        self.payslipButton = tk.Button(top)
        self.payslipButton.place(relx=0.589, rely=0.925, height=30, width=57)
        self.payslipButton.configure(activebackground="#ececec")
        self.payslipButton.configure(activeforeground="#000000")
        self.payslipButton.configure(background="#32C1DC")
        self.payslipButton.configure(disabledforeground="#a3a3a3")
        self.payslipButton.configure(foreground="#000000")
        self.payslipButton.configure(highlightbackground="#32C1DC")
        self.payslipButton.configure(highlightcolor="black")
        self.payslipButton.configure(pady="0")
        self.payslipButton.configure(text='''Payslip''')
        self.payslipButton.configure(command=lambda:self.payslip())

        ####################database buttons start###################

        self.AppointmentTableButton = tk.Button(self.top)
        self.AppointmentTableButton.place(relx=0.012, rely=0.275, height=24, width=183)
        self.AppointmentTableButton.configure(activebackground="#ececec")
        self.AppointmentTableButton.configure(activeforeground="#000000")
        self.AppointmentTableButton.configure(background="#33C6DE")
        self.AppointmentTableButton.configure(disabledforeground="#a3a3a3")
        self.AppointmentTableButton.configure(foreground="#000000")
        self.AppointmentTableButton.configure(highlightbackground="#33C6DE")
        self.AppointmentTableButton.configure(highlightcolor="black")
        self.AppointmentTableButton.configure(pady="0")
        self.AppointmentTableButton.configure(text='''Appointment Table''')
        self.AppointmentTableButton.configure(command=lambda:self.tableButtons(8))

        self.MaterialTableButton = tk.Button(self.top)
        self.MaterialTableButton.place(relx=0.012, rely=0.334, height=24, width=183)
        self.MaterialTableButton.configure(activebackground="#ececec")
        self.MaterialTableButton.configure(activeforeground="#000000")
        self.MaterialTableButton.configure(background="#33C6DE")
        self.MaterialTableButton.configure(disabledforeground="#a3a3a3")
        self.MaterialTableButton.configure(foreground="#000000")
        self.MaterialTableButton.configure(highlightbackground="#33C6DE")
        self.MaterialTableButton.configure(highlightcolor="black")
        self.MaterialTableButton.configure(pady="0")
        self.MaterialTableButton.configure(text='''Material Table''')
        self.MaterialTableButton.configure(command=lambda:self.tableButtons(0))

        self.StaffTableButton = tk.Button(self.top)
        self.StaffTableButton.place(relx=0.012, rely=0.452, height=24, width=183)
        self.StaffTableButton.configure(activebackground="#ececec")
        self.StaffTableButton.configure(activeforeground="#000000")
        self.StaffTableButton.configure(background="#33C6DE")
        self.StaffTableButton.configure(disabledforeground="#a3a3a3")
        self.StaffTableButton.configure(foreground="#000000")
        self.StaffTableButton.configure(highlightbackground="#33C6DE")
        self.StaffTableButton.configure(highlightcolor="black")
        self.StaffTableButton.configure(pady="0")
        self.StaffTableButton.configure(text='''Staff Table''')
        self.StaffTableButton.configure(command=lambda:self.tableButtons(11))

        self.SupplierTableButton = tk.Button(self.top)
        self.SupplierTableButton.place(relx=0.012, rely=0.511, height=24, width=183)
        self.SupplierTableButton.configure(activebackground="#ececec")
        self.SupplierTableButton.configure(activeforeground="#000000")
        self.SupplierTableButton.configure(background="#33C6DE")
        self.SupplierTableButton.configure(disabledforeground="#a3a3a3")
        self.SupplierTableButton.configure(foreground="#000000")
        self.SupplierTableButton.configure(highlightbackground="#33C6DE")
        self.SupplierTableButton.configure(highlightcolor="black")
        self.SupplierTableButton.configure(pady="0")
        self.SupplierTableButton.configure(text='''Supplier Table''')
        self.SupplierTableButton.configure(command=lambda:self.tableButtons(2))

        self.AccountancyTableButton = tk.Button(self.top)
        self.AccountancyTableButton.place(relx=0.012, rely=0.57, height=24, width=183)
        self.AccountancyTableButton.configure(activebackground="#ececec")
        self.AccountancyTableButton.configure(activeforeground="#000000")
        self.AccountancyTableButton.configure(background="#33C6DE")
        self.AccountancyTableButton.configure(disabledforeground="#a3a3a3")
        self.AccountancyTableButton.configure(foreground="#000000")
        self.AccountancyTableButton.configure(highlightbackground="#33C6DE")
        self.AccountancyTableButton.configure(highlightcolor="black")
        self.AccountancyTableButton.configure(pady="0")
        self.AccountancyTableButton.configure(text='''Accountancy Table''')
        self.AccountancyTableButton.configure(command=lambda:self.tableButtons(7))

        self.AppointmentMaterialTableButton = tk.Button(self.top)
        self.AppointmentMaterialTableButton.place(relx=0.012, rely=0.629, height=24, width=183)
        self.AppointmentMaterialTableButton.configure(activebackground="#ececec")
        self.AppointmentMaterialTableButton.configure(activeforeground="#000000")
        self.AppointmentMaterialTableButton.configure(background="#33C6DE")
        self.AppointmentMaterialTableButton.configure(disabledforeground="#a3a3a3")
        self.AppointmentMaterialTableButton.configure(foreground="#000000")
        self.AppointmentMaterialTableButton.configure(highlightbackground="#33C6DE")
        self.AppointmentMaterialTableButton.configure(highlightcolor="black")
        self.AppointmentMaterialTableButton.configure(pady="0")
        self.AppointmentMaterialTableButton.configure(text='''AppointmentMaterial Table''')
        self.AppointmentMaterialTableButton.configure(command=lambda:self.tableButtons(9))

        self.AppointmentAccountancyTableButton = tk.Button(self.top)
        self.AppointmentAccountancyTableButton.place(relx=0.012, rely=0.688, height=24, width=183)
        self.AppointmentAccountancyTableButton.configure(activebackground="#ececec")
        self.AppointmentAccountancyTableButton.configure(activeforeground="#000000")
        self.AppointmentAccountancyTableButton.configure(background="#33C6DE")
        self.AppointmentAccountancyTableButton.configure(disabledforeground="#a3a3a3")
        self.AppointmentAccountancyTableButton.configure(foreground="#000000")
        self.AppointmentAccountancyTableButton.configure(highlightbackground="#33C6DE")
        self.AppointmentAccountancyTableButton.configure(highlightcolor="black")
        self.AppointmentAccountancyTableButton.configure(pady="0")
        self.AppointmentAccountancyTableButton.configure(text='''AppointmentAccountancy Table''')
        self.AppointmentAccountancyTableButton.configure(command=lambda:self.tableButtons(3))

        self.AppointmentTreatmentTableButton = tk.Button(self.top)
        self.AppointmentTreatmentTableButton.place(relx=0.012, rely=0.747, height=24, width=183)
        self.AppointmentTreatmentTableButton.configure(activebackground="#ececec")
        self.AppointmentTreatmentTableButton.configure(activeforeground="#000000")
        self.AppointmentTreatmentTableButton.configure(background="#33C6DE")
        self.AppointmentTreatmentTableButton.configure(disabledforeground="#a3a3a3")
        self.AppointmentTreatmentTableButton.configure(foreground="#000000")
        self.AppointmentTreatmentTableButton.configure(highlightbackground="#33C6DE")
        self.AppointmentTreatmentTableButton.configure(highlightcolor="black")
        self.AppointmentTreatmentTableButton.configure(pady="0")
        self.AppointmentTreatmentTableButton.configure(text='''AppointmentTreatment Table''')
        self.AppointmentTreatmentTableButton.configure(command=lambda:self.tableButtons(4))

        self.AppointmentStaffTableButton = tk.Button(self.top)
        self.AppointmentStaffTableButton.place(relx=0.012, rely=0.806, height=24, width=183)
        self.AppointmentStaffTableButton.configure(activebackground="#ececec")
        self.AppointmentStaffTableButton.configure(activeforeground="#000000")
        self.AppointmentStaffTableButton.configure(background="#33C6DE")
        self.AppointmentStaffTableButton.configure(disabledforeground="#a3a3a3")
        self.AppointmentStaffTableButton.configure(foreground="#000000")
        self.AppointmentStaffTableButton.configure(highlightbackground="#33C6DE")
        self.AppointmentStaffTableButton.configure(highlightcolor="black")
        self.AppointmentStaffTableButton.configure(pady="0")
        self.AppointmentStaffTableButton.configure(text='''AppointmentStaff Table''')
        self.AppointmentStaffTableButton.configure(command=lambda:self.tableButtons(5))

        self.AccountancyStaffTableButton = tk.Button(self.top)
        self.AccountancyStaffTableButton.place(relx=0.012, rely=0.864, height=24, width=183)
        self.AccountancyStaffTableButton.configure(activebackground="#ececec")
        self.AccountancyStaffTableButton.configure(activeforeground="#000000")
        self.AccountancyStaffTableButton.configure(background="#33C6DE")
        self.AccountancyStaffTableButton.configure(disabledforeground="#a3a3a3")
        self.AccountancyStaffTableButton.configure(foreground="#000000")
        self.AccountancyStaffTableButton.configure(highlightbackground="#33C6DE")
        self.AccountancyStaffTableButton.configure(highlightcolor="black")
        self.AccountancyStaffTableButton.configure(pady="0")
        self.AccountancyStaffTableButton.configure(text='''AccountancyStaff Table''')
        self.AccountancyStaffTableButton.configure(command=lambda:self.tableButtons(6))

        self.TreatmentTableButton = tk.Button(self.top)
        self.TreatmentTableButton.place(relx=0.012, rely=0.393, height=24, width=183)
        self.TreatmentTableButton.configure(activebackground="#ececec")
        self.TreatmentTableButton.configure(activeforeground="#000000")
        self.TreatmentTableButton.configure(background="#33C6DE")
        self.TreatmentTableButton.configure(disabledforeground="#a3a3a3")
        self.TreatmentTableButton.configure(foreground="#000000")
        self.TreatmentTableButton.configure(highlightbackground="#33C6DE")
        self.TreatmentTableButton.configure(highlightcolor="black")
        self.TreatmentTableButton.configure(pady="0")
        self.TreatmentTableButton.configure(text='''Treatment Table''')
        self.TreatmentTableButton.configure(command=lambda:self.tableButtons(1))

        self.PatientTableButton = tk.Button(self.top)
        self.PatientTableButton.place(relx=0.012, rely=0.923, height=24, width=183)
        self.PatientTableButton.configure(activebackground="#ececec")
        self.PatientTableButton.configure(activeforeground="#000000")
        self.PatientTableButton.configure(background="#33C6DE")
        self.PatientTableButton.configure(disabledforeground="#a3a3a3")
        self.PatientTableButton.configure(foreground="#000000")
        self.PatientTableButton.configure(highlightbackground="#33C6DE")
        self.PatientTableButton.configure(highlightcolor="black")
        self.PatientTableButton.configure(pady="0")
        self.PatientTableButton.configure(text='''Patient Table''')
        self.PatientTableButton.configure(command=lambda:self.tableButtons(10))

        ####################database buttons end###################
    '''
    this function uses recursion to reverse the elements within a list
    '''
    def reverseList(self,l, k):
        if len(l) == 0: return k([])
        def b(res):
            return k([l[-1]] + res)
        return self.reverseList(l[:-1],b)
    '''
    this will check what tables the user has access to
    '''
    def accessCheck(self):
        if self.accessLevel == 0:
            accessGranted = True
        elif self.accessLevel == 1:
            if self.table_name == "Patient":
                accessGranted = False
            else:
                accessGranted = True
        elif self.accessLevel == 2:
            if self.table_name == "Staff" or self.table_name == "Appointment" or self.table_name == "AppointmentMaterial" or self.table_name == "AppointmentStaff" or self.table_name == "AppointmentTreatment" or self.table_name == "Treatment":
                accessGranted = True
            else:
                accessGranted = False
        elif self.accessLevel == 3:
            if self.table_name == "Staff" or self.table_name == "Appointment" or self.table_name == "Patient":
                accessGranted = True
            else:
                accessGranted = False
        elif self.accessLevel == 4:
            if self.table_name == "Staff" or self.table_name == "AppointmentStaff":
                accessGranted = True
            else:
                accessGranted = False
        else:
            totalAccessGranted = False
            messagebox.showerror("Error", "The access level not found please speak to the admin.")
        return accessGranted
    '''
    this creates the drop down menu for search
    '''
    def dropDownMenu(self,table):
        if table == None:
            OptionList = ["Please Select a Table"]
        else:
            self.fields = self.getting_fields_from_a_table()
            OptionList = self.fields
            if self.table_name == "Appointment":
                if self.accessLevel == 0 or self.accessLevel == 1:
                    pass
                else:
                    self.fields.pop(6)
                    self.fields.pop(5)

        self.searchList = tk.StringVar(self.top)
        self.searchList.set(OptionList[0])

        self.searchWidget = tk.OptionMenu(self.top, self.searchList, *OptionList)
        self.searchWidget.place(relx=0.238, rely=0.2, height=24, width=190)
        self.searchWidget.configure(activebackground="#ececec")
        self.searchWidget.configure(activeforeground="#000000")
        self.searchWidget.configure(background="#33C6DE")
        self.searchWidget.configure(disabledforeground="#a3a3a3")
        self.searchWidget.configure(foreground="#000000")
        self.searchWidget.configure(highlightbackground="#ffffff")
        self.searchWidget.configure(highlightcolor="black")
        self.searchWidget.configure(pady="0")

    """
    this will create a calendar widget which will be pre selected with todays date
    """
    def calendar(self):
        self.root = tk.Tk()

        def sel():
            self.selectedDate = cal.selection_get()
            self.selectedDate = self.selectedDate.strftime("%d/%m/%Y")
            self.root.destroy()

        today = datetime.datetime.now()
        cal = Calendar(self.root,font="Arial 14", selectmode='day',cursor="hand1", day = today.day, month = today.month,year = today.year)
        cal.pack(fill="both", expand=True)
        ttk.Button(self.root, text="ok", command=lambda:sel()).pack()
    '''
    This will check if a user has selected a table whenever they press the add, edit, delete or search button.
    It will also check which button has been pressed and if a table is selected it will run the function for
    wahtever button that has been pressed.
    '''
    def checkButtonPressed(self, buttonPressed):
        tabsOpen = 0
        for i in self.nb.tabs():
            tabsOpen = tabsOpen+1
        if tabsOpen != 0:
            if buttonPressed == 1:
                self.add_tab(1)
            if buttonPressed == 2:
                self.add_tab(2)
            if buttonPressed == 3:
                self.delete_record()
            if buttonPressed == 4:
                self.search(self.SearchEntry)
        else:
            messagebox.showerror("Error", "Please select a table first.")
    '''
    Below is the function which is responsible for the creation and sending of the payslips.
    '''
    def payslip(self):
        if self.accessLevel == 0 or self.accessLevel == 1:
            if messagebox.askyesno("Email", "Are you sure you want to send out payslips to all staff members? This will take a long time and the database will be out of operation while the emails are being sent."):
                try:
                    conn = sqlite3.connect(self.db_name)
                    cur = conn.cursor()
                    cur.execute("SELECT StaffID FROM Staff")
                    staffID = cur.fetchall()
                    staffID = self.unpackList(staffID)
                    today = datetime.datetime.now()
                    Today = today.strftime("%d/%m/%Y")
                    date = today - datetime.timedelta(days=7)
                    date = date.strftime("%d/%m/%Y")
                    timePeriod = date + " - " + Today
                    for i in staffID:
                        cur.execute("SELECT Email FROM Staff WHERE StaffID = {0}".format(i))
                        staffEmail, = cur.fetchone()
                        cur.execute("SELECT Name FROM Staff WHERE StaffID = {0}".format(i))
                        staffName, = cur.fetchone()
                        cur.execute("SELECT WeeklyPay FROM Staff WHERE StaffID = {0}".format(i))
                        weeklyPay, = cur.fetchone()
                        
                        password = "yesemail"

                        msg = MIMEMultipart()
                        msg['From'] = "smtplib66@gmail.com"
                        msg['To'] = "{0}".format(staffEmail)
                        msg['Subject'] = "Payslip"

                        body = "Hello {0}, The following is the information for your payslip this week. Period worked: {1}, Pay for this period: {2}. From SMILEXE.".format(staffName,timePeriod,weeklyPay)
                        msg.attach(MIMEText(body,'html')) #This puts the body of the message into a html format

                        server = smtplib.SMTP("smtp.gmail.com", 587)
                        server.starttls() #encrypts login
                        server.login(msg['From'], password)
                        server.sendmail(msg['From'], msg['To'], msg.as_string())
                        server.quit()
                    conn.close()
                    messagebox.showinfo("Emails", "All payslips have been sent")
                except smtplib.SMTPRecipientsRefused:
                    messagebox.showerror("Error","The email address you are trying to send to does not exist.")
        else:
            messagebox.showerror("Error","You do not have the correct permissions for this task.")

    '''
    This is responsible for the creation and sending of the patients reminder email.
    '''
    def reminder_email(self, time,date,patientID):#needs patientID to find what email to send to

        if messagebox.askyesno("Email", "Are you sure you want to send a reminder email?"):
            try:
                if patientID == "":
                    messagebox.showerror("Error", "Please select a PatientID to send a reminder email")
                else:
                    conn = sqlite3.connect(self.db_name)
                    cur = conn.cursor()
                    cur.execute("SELECT Email FROM Patient WHERE PatientID = {0}".format(patientID))
                    patients_email, = cur.fetchone()
                    cur.execute("SELECT Name FROM Patient WHERE PatientID = {0}".format(patientID))
                    patientsName, = cur.fetchone()
                    conn.close()

                    password = "yesemail"

                    msg = MIMEMultipart()
                    msg['From'] = "smtplib66@gmail.com"
                    msg['To'] = "{0}".format(patients_email)
                    msg['Subject'] = "Dentist Reminder"

                    body = """Dear {0},
                            This is just a confirmation of your dentist appointment. The appointment you have booked is at {1} on the {2}. If there is any problem with this time or date please contact the dentist office as soon as possible.
                            Thank you.
                            From SMILEXE.""".format(patientsName,time,date)
                    msg.attach(MIMEText(body,'html')) #This puts the body of the message into a html format

                    server = smtplib.SMTP("smtp.gmail.com", 587)
                    server.starttls() #encrypts login
                    server.login(msg['From'], password)
                    server.sendmail(msg['From'], msg['To'], msg.as_string())
                    server.quit()
                    messagebox.showinfo("Emails", "Reminder email has been sent.")
            except smtplib.SMTPRecipientsRefused:
                messagebox.showerror("Error","The email address you are trying to send to does not exist.")
    '''
    This is responsible for the creation of the report and the saving of the report.
    '''
    def report(self):
        if self.accessLevel == 0 or self.accessLevel == 1:
            today = datetime.datetime.now()
            Today = today.strftime("%d/%m/%Y")

            self.graph()

            rows=[]
            
            document = Document()

            document.add_heading('Report', 0)

            p = document.add_paragraph('This is a report generated on the ')
            p.add_run('{0}.'.format(Today)).bold = True

            document.add_heading('Staffs Pay', level=1)

            conn = sqlite3.connect(self.db_name)
            cur = conn.cursor()
            cur.execute("Select Name from Staff")
            names = cur.fetchall()
            cur.execute("select WeeklyPay from Staff")
            weeklyPayList = cur.fetchall()
            names = self.unpackList(names)
            weeklyPayList = self.unpackList(weeklyPayList)

            for i in range(len(names)):
                name = names[i]
                pay = weeklyPayList[i]
                document.add_paragraph('Staff Name: {0}, Weekly Pay: {1}'.format(name,pay), style='List Bullet')

            document.add_heading('Graph of the profit for the past month:', level=1)

            document.add_paragraph('Data taken from the accountancy table.')

            document.add_picture('figure1.png', width=Inches(6))

            saveFilePath  =  filedialog.asksaveasfilename(title = "Save Report",defaultextension=".docx",filetypes = (("docx files","*.docx"),("all files","*.*")))

            try:
                document.save(saveFilePath)
            except FileNotFoundError:
                messagebox.showerror("Error","You did not save the report.")
        else:
            messagebox.showerror("Error", "You do not have access to perform this task.")
    '''
    This is responsible for the creation of the graph and saving of the graph for the report.
    '''
    def graph(self):

        conn = sqlite3.connect(self.db_name)
        cur = conn.cursor()
        today = datetime.datetime.now()
        Today = today.strftime("%d/%m/%Y")
        values = []
        datesInRange = []
        for i in range(28): #this for loop will check the accountancy record for any dates within the past month and then will add the dates that has happened in the past month to the graph
            date = today - datetime.timedelta(days=i)
            date = date.strftime("%d/%m/%Y")
            speechmark = "'"
            dateWithSpeechmark = speechmark + date + speechmark
            cur.execute("SELECT Profit FROM Accountancy WHERE DateOfEntry == {0}".format(dateWithSpeechmark))
            dataCollected = cur.fetchall()
            dataCollected = self.unpackList(dataCollected)
            if dataCollected != []:
                date = parser.parse(date)
                dateFrom = date - datetime.timedelta(days=7)
                date = date.strftime("%d/%m/%Y")
                dateFrom = dateFrom.strftime("%d/%m/%Y")
                date = str(dateFrom) + " - " + str(date)
                values.append(dataCollected[0])
                datesInRange.append(date)
                
        datesInRange = self.reverseList(datesInRange, lambda x:x)
        values = self.reverseList(values, lambda x:x)

        plt.figure(figsize=(10, 5))
        plt.bar(datesInRange, values, color = "red")
        plt.xlabel("Profit")
        plt.ylabel("Date recorded")
        plt.suptitle('Report Grpah')
        plt.savefig('figure1.png')
        plt.close()
        
    '''
    the function below will query the database
    '''
    def query(self, sql,data):
        with sqlite3.connect(self.db_name) as db:
            cursor = db.cursor()
            cursor.execute("PRAGMA foreign_keys = OFF")
            cursor.execute(sql,data)
            db.commit()
    '''
    the function below will retrieve the record selected within the treeview and return it within a
    2d array
    '''
    def get_record(self): # This will return the data in a selected record.
        focused=self.tview.selection()
        records=[]
        if focused != "": # This prevents an error if nothing has been selected.
            for i in focused:
                index_value=str(i)[1:]# This disreagrds the first character of a string.
                index_value=int(index_value,16)# This changes the hex number into decimal.
                for child in self.tview.get_children():
                    x=str(child)[1:]
                    x=int(x,16)
                    if x==index_value:
                        record = self.tview.item(child)["values"]
                records.append(record)
            return records
    '''
    this function will refresh the treeview when a change is made to the database
    '''
    def refresh(self):
        try:
            for row in self.tview.get_children():
                self.tview.delete(row)
            conn = sqlite3.connect(self.db_name)
            cur = conn.cursor()
            cur.execute("SELECT * FROM {}".format(self.table_name))
            rows = cur.fetchall()
            for row in rows:
                self.tview.insert("", tk.END, values=row)
            if self.table_name == "Staff" or self.table_name == "AppointmentStaff":
                if self.accessLevel == 0 or self.accessLevel == 1:
                    pass
                else:
                    self.tview.delete(*self.tview.get_children())
                    cur.execute("SELECT * FROM {0} where StaffID = {1}".format(self.table_name,self.StaffID))
                    records = cur.fetchall()
                    for record in records:
                        self.tview.insert("", tk.END, values=record)
            conn.close()
        except AttributeError:
            messagebox.showerror("Error","Please select a table first.")
    '''
    this function will delete a record within the database which has been selected within the treeview 
    '''
    def delete_record(self):
        accessGranted = True
        if self.table_name == "Staff" or self.table_name == "AppointmentStaff":
            if self.accessLevel == 0 or self.accessLevel == 1:
                accessGranted = True
            else:
                accessGranted = False
        if accessGranted == True:
            records_to_delete = self.get_record()
            if records_to_delete == []:
                messagebox.showerror("Error", "Please select a record to delete.")
            else:
                if messagebox.askyesno("Delete", "Are you sure you want to delete this record: {0}".format(records_to_delete)):
                    self.fields = self.getting_fields_from_a_table()
                    pk = self.fields[0]
                    data = records_to_delete[0]
                    data = str(data[0])
                    if self.table_name == "AppointmentMaterial" or self.table_name == "AppointmentTreatment":
                        self.undoingProcessing(records_to_delete)
                    with sqlite3.connect(self.db_name) as db:
                        cursor = db.cursor()
                        sql = "delete from {0} where {1}={2}".format(self.table_name,pk,data)
                        cursor.execute(sql)
                        db.commit()
                    self.refresh()
        else:
            messagebox.showwarning("Warning","You do not have the correct access level to perform this task.")
    '''
    this function will be responsible for searching for a record in the database
    '''
    def search(self, SearchEntry):
        searchterm = str(self.SearchEntry.get())
        if searchterm == "":
            self.refresh()
        else:
            selection = self.searchList.get()
            searchterm = "'" + '%'+ searchterm +'%' + "'"
            conn = sqlite3.connect(self.db_name)
            cur = conn.cursor()
            cur.execute("SELECT * FROM {0} WHERE {1} LIKE {2}".format(self.table_name,selection,searchterm))
            records  = cur.fetchall()
            if records != []:
                self.tview.delete(*self.tview.get_children())
                for record in records:
                    self.tview.insert('', 'end', values=(record))
                if self.table_name == "Staff" or self.table_name == "AppointmentStaff":
                    if self.accessLevel == 0 or self.accessLevel == 1:
                        pass
                    else:
                        self.refresh()
                        messagebox.showwarning("Warning","You do not have the correct access level to perform this task.")
            else:
                messagebox.showwarning("Results","Not Found")

    '''
    the function below will get the names of all the tables within the database and return it
    as an array
    '''
    def getting_tables(self):
        with sqlite3.connect(self.db_name) as db:
            cur = db.cursor()
            cur.execute("select name from sqlite_master")
            names=cur.fetchall()
            return names
    '''
    this function will return all the self.fields within a table as an array
    '''
    def getting_fields_from_a_table(self):
        field_names_of_table=[]
        with sqlite3.connect(self.db_name) as db:
            cur = db.cursor()
            cur.execute("PRAGMA TABLE_INFO ({0})".format(self.table_name)) #This gets the field names, the data types and other stuff
            list_of_fields=cur.fetchall() 
            for i in list_of_fields:
                field_name=i[1] #The reason for this is because the field names are stored in element 1 in ecah tuple in the list. 
                field_names_of_table.append(field_name)
            return field_names_of_table
        
    '''
    this creates a dropdown menu for the adding or editing of records that need list validation
    '''
    def dropDownForRecords(self,options,frame):
        dropDownSelection = tk.StringVar(frame)
        dropDownSelection.set(options[0])
        dropDownButton =  tk.OptionMenu(frame, dropDownSelection, *options)
        dropDownButton.grid(row=1,column=1)
        def callback(*args):
            selectionMade = dropDownSelection.get()
            return selectionMade
        return dropDownSelection
    '''
    the add_tab() function will create the tab within the notebook which will have a label and entry widget for every
    field within the table which is currently open. 
    '''
    def add_tab(self,check_button_pressed):
        accessGranted = True
        if self.table_name == "Staff" or self.table_name == "AppointmentStaff":
            if self.accessLevel == 0 or self.accessLevel == 1:
                accessGranted = True
            else:
                accessGranted = False
        if accessGranted == True:
            self.fields = self.getting_fields_from_a_table()
            new_frame = tk.Frame(self.nb, background="#ffffff")
            count = 0
            widgets = []
            self.fields.pop(0)#removes the primary key from the self.fields the user can add
            if self.table_name == "Staff":
                self.fields.pop(10)#removes WeeklyPay as it is generated by the system
            if self.table_name == "Accountancy":
                for i in range(2):
                    self.fields.pop(1)
            if self.table_name == "Appointment":
                self.fields.pop(5)
                self.fields.pop(4)
            for i in self.fields:
                count = count+1
                tk.Label(new_frame,background="#ffffff",text="{}:".format(i)).grid(row=count,column=0)
                entry = tk.Entry(new_frame)
                widgets.append(entry)
                entry.grid(row=count,column=1)
            callback = None
            dropDownSelection = None
            if self.table_name == "Appointment":
                calButton = tk.Button(new_frame,text="            Calendar            ",command=lambda:self.calendar())
                warningLabel1 = tk.Label(new_frame,background="#ffffff",text="TotalPrice and TotalProfit For Appointment")
                warningLabel2 = tk.Label(new_frame,background="#ffffff",text="wil be updated based on values entered into")
                warningLabel3 = tk.Label(new_frame,background="#ffffff",text="AppointmentMaterial and AppointmentTreatment.")
                warningLabel1.grid(row=1,column=2)
                warningLabel2.grid(row=2,column=2)
                warningLabel3.grid(row=3,column=2)
                calButton.grid(row=2,column=1)
                timeSlots = ["    09:00-10:00    ", "    10:00-11:00    ", "    11:00-12:00    ", "    12:00-14:00    ", "    14:00-16:00    ", "    16:00-17:00    "]
                dropDownSelection = self.dropDownForRecords(timeSlots,new_frame)
            if self.table_name == "Staff":
                listOfAccessLevels = ["             0             ","             1            ","             2             ","             3             ","             4             "]
                dropDownSelection = self.dropDownForRecords(listOfAccessLevels,new_frame)
            if check_button_pressed == 1:
                button = tk.Button(new_frame,text="Add record",command=lambda:self.insert_data(widgets,new_frame,callback,dropDownSelection,check_button_pressed))
                self.nb.add(new_frame,text="Add")
            if check_button_pressed == 2:
                selectedRecord = self.get_record()
                selectedRecord = selectedRecord[0]
                selectedRecord.pop(0)
                for i in range(len(widgets)):
                    widgets[i].insert(10,selectedRecord[i])
                button = tk.Button(new_frame,text="Edit record",command=lambda:self.edit_data(widgets,new_frame,callback,dropDownSelection,check_button_pressed,selectedRecord))
                self.nb.add(new_frame,text="Edit")
                
            self.nb.select(new_frame)
            button.grid(column=0,pady=4,padx=4)

        else:
            messagebox.showwarning("Warning","You do not have the correct access level to perform this task.")
    '''
    This is responsible for all calculations that need to be done which have no effect on other tables
    '''
    def processing(self,widgets,new_frame,callback,dropDownSelection,checkButtonPressed,selectedRecord):
        records = []
        for i in range(len(widgets)):
            record = widgets[i]
            records.append(record.get())

        if self.table_name == "Appointment" or self.table_name == "Staff":
            dropDownSelection.trace("w", callback)
            selectionMade = dropDownSelection.get().strip()
            records[0] = selectionMade
            
        if self.table_name == "Appointment":
            try:
                records[1] = self.selectedDate
            except AttributeError:
                records[1] = " "
                if checkButtonPressed == 2:
                    records[1] = selectedRecord[1]
            

            if checkButtonPressed == 2:
                records.append(selectedRecord[4])
                records.append(selectedRecord[5])
            else:
                records.append("0.0")
                records.append("0.0")
                
            self.fields.append("TotalPrice")
            self.fields.append("TotalProfitForAppointment")
                
            self.reminder_email(records[0],records[1],records[3])
            
        if self.table_name == "Staff":
            self.presenceCheck(records[1],"password") #makes sure something is entered for password
            self.typeCheck(records[2],"str","Name") #makes sure Name is entered as a string value and not a float
            self.presenceCheck(records[3],"email") #makes sure something is entered for email
            self.lengthCheck(records[4],11,11,"PhoneNumber") #checks that PhoneNumber is within the length 11-11
            self.presenceCheck(records[5],"Address") #makes sure something is entered for address
            self.formatCheck(records[6],"[A-Z]{1,2}[0-9][A-Z0-9]{1}? [0-9][ASD-HJLNP-UW-Z]{2}","Postcode","LLNN NLL") #makes sure postcode is within the format of LLNN NLL
            self.presenceCheck(records[7],"JobTitle") #makes sure something is entered for JobTitle
            self.rangeCheck(records[8],40,0,"HoursWorkedPerWeek") #checks that HoursWorkedPerWeek is within the range 0-40
            if self.error == False:
                try:
                    weeklyPay = float(records[8])*float(records[9])
                    records.append(str(weeklyPay))
                    self.fields.append("WeeklyPay")
                except ValueError:
                    messagebox.showerror("Error","Please enter float values for PayPerHour and HoursWorkedPerWeek.")
                    self.error = True

        if self.table_name == "Accountancy":
            if checkButtonPressed == 2:
                try:
                    profit = (float(selectedRecord[1]) + float(selectedRecord[0])) - float(records[0])
                    records.append(str(profit))
                    records.append(str(selectedRecord[2]))
                    self.fields.append("Profit")
                    self.fields.append("DateOfEntry")
                except ValueError:
                    messagebox.showerror("Error","Please enter a float value for OtherExpenses")
            else:
                today = datetime.datetime.now()
                Today = today.strftime("%d/%m/%Y")
                listOfAllProfitFromAppointments = []
                conn = sqlite3.connect(self.db_name)
                cur = conn.cursor()
                weekAgo = today - datetime.timedelta(days=7)
                weekAgo = weekAgo.strftime("%d/%m/%Y")
                weekAgo = "'" + weekAgo + "'"
                cur.execute("SELECT * FROM Accountancy WHERE DateOfEntry == {0}".format(weekAgo))
                found = cur.fetchone()
                if found == None:
                    messagebox.showerror("Error","This is not the correct date to add an accountancy record this must be done a week apart exaclty.")
                    self.error = True
                else:
                    for i in range(7):
                        date = today - datetime.timedelta(days=i)
                        date = date.strftime("%d/%m/%Y")
                        date = "'" + date + "'"
                        cur.execute("SELECT TotalProfitForAppointment FROM Appointment WHERE DateOfAppointment == {0}".format(date))
                        dataCollected = cur.fetchall()
                        listOfAllProfitFromAppointments.append(dataCollected)
                    totalProfit = []
                    for i in range(len(listOfAllProfitFromAppointments)):
                        tempList = listOfAllProfitFromAppointments[i]
                        for j in range(len(tempList)):
                            totalProfitOfAnAppoitment, = tempList[j]
                            float(totalProfitOfAnAppoitment)
                            totalProfit.append(totalProfitOfAnAppoitment)
                    sumOfAppointmentProfit = sum(totalProfit)
                    cur.execute("SELECT WeeklyPay FROM Staff")
                    weeklyPay = cur.fetchall()
                    listOfWeeklyPay = []
                    for i in range(len(weeklyPay)):
                        employeePay, = weeklyPay[i]
                        float(employeePay)
                        listOfWeeklyPay.append(employeePay)
                    sumOfWeeklyPay = sum(listOfWeeklyPay)
                    try:
                        OtherExpenses = float(records[0])
                        profit = sumOfAppointmentProfit - sumOfWeeklyPay - OtherExpenses
                        records.append(str(profit))
                        records.append(str(Today))
                        self.fields.append("Profit")
                        self.fields.append("DateOfEntry")
                    except ValueError:
                        messagebox.showerror("Error","Please enter a float value for OtherExpenses")
            
        return records

    '''
    The following if statements are for processing data.
    They are here and not in the processing subroutine because this requires changes to another table so the user needs to confirm that the changes that will be made are correct before processing this data.
    '''

    def undoingProcessing(self, selectedRecord):
        selectedRecord = selectedRecord[0]
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        if self.table_name == "AppointmentMaterial":    
            cursor.execute("select CurrentStock from Material where MaterialID = {0}".format(selectedRecord[2]))
            selectedCurrentStock, = cursor.fetchone()
            cursor.execute("select Price from Material where MaterialID = {0}".format(selectedRecord[2]))
            selectedPriceOfMaterial, = cursor.fetchone()
            selectedAmountOfMaterialsUsed = selectedRecord[3]
            updatedCurrentStock = int(selectedCurrentStock) + int(selectedAmountOfMaterialsUsed)
            tempList = []
            tempList.append(updatedCurrentStock)
            sql = ("update Material set CurrentStock=? where MaterialID = {0}".format(selectedRecord[2]))
            cursor.execute(sql,tempList)
            cursor.execute("select TotalProfitForAppointment from Appointment where AppointmentID = {0}".format(selectedRecord[1]))
            selectedTotalProfitForAppointment, = cursor.fetchone()
            selectedTotalPriceOfMaterialsUsed = float(selectedAmountOfMaterialsUsed) * float(selectedPriceOfMaterial)
            originalTotalProfitForAppointment = float(selectedTotalProfitForAppointment) + float(selectedTotalPriceOfMaterialsUsed)
            sql = ("update Appointment set TotalProfitForAppointment=? where AppointmentID = {0}".format(selectedRecord[1]))
            tempList = []
            tempList.append(originalTotalProfitForAppointment)
            cursor.execute(sql,tempList)
            conn.commit()
        if self.table_name == "AppointmentTreatment":
            cursor.execute("select Price from Treatment where TreatmentID = {0}".format(selectedRecord[1]))
            selectedPrice, = cursor.fetchone()
            cursor.execute("select TotalProfitForAppointment from Appointment where AppointmentID = {0}".format(selectedRecord[2]))
            selectedTotalProfitForAppointment, = cursor.fetchone()
            cursor.execute("select TotalPrice from Appointment where AppointmentID = {0}".format(selectedRecord[2]))
            selectedTotalPrice, = cursor.fetchone()
            selectedTotalProfitForAppointment = float(selectedTotalProfitForAppointment) - float(selectedPrice)
            selectedTotalPrice = float(selectedTotalPrice) - float(selectedPrice)
            tempList = []
            tempList.append(str(selectedTotalPrice))
            sql =("update Appointment set TotalPrice=? where AppointmentID = {0}".format(selectedRecord[2]))
            cursor.execute(sql,tempList)
            tempList = []
            tempList.append(str(selectedTotalProfitForAppointment))
            sql = ("update Appointment set TotalProfitForAppointment=? where AppointmentID = {0}".format(selectedRecord[2]))
            cursor.execute(sql,tempList)
            conn.commit()
    '''
    This is responsible for all calculations that need to be done which have an effect on other tables
    '''
    def processing2(self,records,check_button_pressed):
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        try:
            if self.table_name == "AppointmentMaterial":
                if check_button_pressed == 2:
                    selectedRecord = self.get_record(self.tview)
                    self.undoingProcessing(selectedRecord)
                amountOfMaterialUsed = records[2]
                sql = "select CurrentStock from Material where MaterialID = {0}".format(records[1])
                cursor.execute("select CurrentStock from Material where MaterialID = {0}".format(records[1]))
                currentStock, = cursor.fetchone()
                newStock = int(currentStock)-int(amountOfMaterialUsed)
                cursor.execute("select Price from Material where MaterialID = {0}".format(records[1]))
                priceOfMaterial, = cursor.fetchone()
                totalPriceOfMaterialsUsed = float(amountOfMaterialUsed) * float(priceOfMaterial)
                cursor.execute("select TotalProfitForAppointment from Appointment where AppointmentID = {0}".format(records[0]))
                TotalProfitForAppointment, = cursor.fetchone()
                TotalProfitForAppointment = float(TotalProfitForAppointment) - float(totalPriceOfMaterialsUsed)
                str(TotalProfitForAppointment)
                tempList = []
                tempList.append(newStock)
                sql = ("update Material set CurrentStock=? where MaterialID = {0}".format(records[1]))
                cursor.execute(sql,tempList)
                sql = ("update Appointment set TotalProfitForAppointment=? where AppointmentID = {0}".format(records[0]))
                tempList = []
                tempList.append(TotalProfitForAppointment)
                cursor.execute(sql,tempList)
                conn.commit()
                if newStock ==20 or newStock < 20:
                    messagebox.showwarning("Warning", "Stock of this material is running low.")
            if self.table_name == "AppointmentTreatment":
                if check_button_pressed == 2:
                    selectedRecord = self.get_record(self.tview)
                    self.undoingProcessing(selectedRecord)
                cursor.execute("select Price from Treatment where TreatmentID = {0}".format(records[0]))
                price, = cursor.fetchone()
                cursor.execute("select TotalPrice from Appointment where AppointmentID = {0}".format(records[1]))
                totalPrice, = cursor.fetchone()
                totalPrice = float(totalPrice) + float(price)
                tempList = []
                tempList.append(str(totalPrice))
                sql =("update Appointment set TotalPrice=? where AppointmentID = {0}".format(records[1]))
                cursor.execute(sql,tempList)
                cursor.execute("select TotalProfitForAppointment from Appointment where AppointmentID = {0}".format(records[1]))
                TotalProfitForAppointment, = cursor.fetchone()
                TotalProfitForAppointment = float(TotalProfitForAppointment) + float(price)
                sql = ("update Appointment set TotalProfitForAppointment=? where AppointmentID = {0}".format(records[1]))
                tempList = []
                tempList.append(TotalProfitForAppointment)
                cursor.execute(sql,tempList)
                conn.commit()
        except TypeError:
            messagebox.showerror("Error","One or both of the ID's you entered do not exist yet.")
            self.error = True
        conn.close()

    '''
    this function will be responsible for adding the data entered by the user into the database
    '''
    def insert_data(self,widgets,new_frame,callback,dropDownSelection,check_button_pressed):

        self.error = False
        records = self.processing(widgets,new_frame,callback,dropDownSelection,check_button_pressed,None)
        nameOfFields = ','.join(self.fields)
        if self.error == False:
            if messagebox.askyesno("Add", "Are you sure you want to add this record: {0}".format(records)):
                values = []
                for i in records:
                    values.append("?")
                values = ','.join(values)
                self.processing2(records,check_button_pressed)
                if self.error == False:
                    sql = "insert into {0}({1}) values ({2})".format(self.table_name,nameOfFields,values)
                    self.query(sql,records)
                    self.nb.forget(new_frame)
                    self.refresh()
                else:
                    self.error = False
        else:
            self.error = False
    '''
    the function below will edit the selected record to the data which was entered into the edit tab
    '''
    
    def edit_data(self,widgets,new_frame,callback,dropDownSelection,check_button_pressed,selectedRecord):

        self.error = False

        records = self.processing(widgets,new_frame,callback,dropDownSelection,check_button_pressed,selectedRecord)

        edit_fields = []
        records_to_edit = self.get_record()
        if self.error == False:
            if records_to_edit == []:
                messagebox.showerror("Error","Please select a record to edit.")
            else:
                if messagebox.askyesno("Edit", "Are you sure you want to edit the record {0} to this {1}".format(records_to_edit,records)):
                    getting_pk = self.getting_fields_from_a_table()
                    pk = getting_pk[0]
                    data = records_to_edit[0]
                    data = str(data[0])
                    self.processing2(records,check_button_pressed)
                    if self.error == False:
                        for i in range(len(records)):
                            work = str(self.fields[i]) + "=?"
                            edit_fields.append(work)
                        nameOfFields = ','.join(edit_fields)
                        self.nb.forget(new_frame)
                        with sqlite3.connect(self.db_name) as db:
                            cursor = db.cursor()
                            sql = "update {0} set {1} where {2}={3}".format(self.table_name,nameOfFields,pk,data)
                            cursor.execute(sql,records)
                            db.commit()
                            self.refresh()
                    else:
                        self.error = False
        else:
            self.error = False
    '''
    the function below will be creating the notebook and the treeview which the table will be displayed in
    '''
    def Create(self,table):

        frame=tk.Frame(self.nb, background="#ffffff")
        self.table_names = self.getting_tables()
        self.currentTableName, = self.table_names[table]

        self.nb.add(frame, text="{}".format(self.table_name)) #Adds the tab.
        self.nb.pack(expand=1,fill="both")

        #creating the self.tview
        self.tview=ttk.Treeview(frame, selectmode='browse', show="headings")
        self.fields = self.getting_fields_from_a_table()
        columns = []
        for i in range(len(self.fields)):
            slot = "SLOT_{0}".format(i)
            columns.append(slot)
            self.tview["columns"]=(columns)
            self.tview["show"]="headings" #This hides column 0.
            self.tview.column(slot)
        for i in range(len(columns)):
            self.tview.heading("SLOT_{0}".format(i),text=self.fields[i])
        for row in self.tview.get_children():
            self.tview.delete(row)
        conn = sqlite3.connect(self.db_name)
        cur = conn.cursor()
        cur.execute("SELECT * FROM {}".format(self.table_name))
        rows = cur.fetchall()
        for row in rows:
            self.tview.insert("", tk.END, values=row)
        if self.table_name == "Staff" or self.table_name == "AppointmentStaff":
            if self.accessLevel == 0 or self.accessLevel == 1:
                pass
            else:
                self.tview.delete(*self.tview.get_children())
                cur.execute("SELECT * FROM {0} where StaffID = {1}".format(self.table_name,self.StaffID))
                records = cur.fetchall()
                for record in records:
                    self.tview.insert("", tk.END, values=record)
        if self.table_name == "Appointment":
            if self.accessLevel == 0 or self.accessLevel == 1 or self.accessLevel == 4:
                pass
            else:
                self.tview["displaycolumns"]=("0","1","2","3")

        conn.close()

        #creating both the horizontal and vertical scroll bar
        vsb=ttk.Scrollbar(frame, orient="vertical", command=self.tview.yview)
        vsb.pack(side="right",fill="y") 

        hsb=ttk.Scrollbar(frame, orient="horizontal", command=self.tview.xview)
        hsb.pack(side="bottom",fill="x")

        self.tview.configure(yscrollcommand=vsb.set)
        self.tview.configure(xscrollcommand=hsb.set)

        self.dropDownMenu(self.table_name)
                
        self.tview.pack(expand=True,fill="both") #This will expand to the area in the frame.

    '''
    this function will check if a tab is open and if a tab is open it will close all the tabs
    '''
    def check(self):
        for i in self.nb.tabs():
            self.nb.forget(i)
    '''
    this is responsible for telling Create what table to create based on the button pressed 
    '''
    def tableButtons(self, table):
        self.table_names = self.getting_tables()
        self.table_name, = self.table_names[table]
        access = self.accessCheck()
        if access == True:
            self.check()
            self.Create(table)
        else:
            self.table_name = self.currentTableName
            messagebox.showwarning("Warning","You do not have access to this table.")

if __name__ == "__main__":
    Login(True)
